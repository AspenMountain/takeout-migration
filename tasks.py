"""
tasks.py
--------
Parse Google Tasks from a Takeout export and render to DOCX.

Observed Takeout schema (as of 2026)
-------------------------------------
Single file  Tasks/Tasks.json  with:
    kind  : "tasks#taskLists"
    items : list of task-list dicts, each containing:
        kind    : "tasks#tasks"
        id      : str
        title   : str
        updated : ISO-8601 timestamp
        items   : list of task dicts

Each task dict:
    id        : str
    title     : str
    status    : "needsAction" | "completed"
    notes     : str             (optional)
    due       : ISO-8601 timestamp  (optional)
    completed : ISO-8601 timestamp  (optional, when status=="completed")
    created   : ISO-8601 timestamp
    parent    : str             (id of parent task; absent on root tasks)
    starred   : bool            (optional)
"""

from __future__ import annotations

import datetime as dt
import io
import json
import os
from pathlib import Path


# ---------------------------------------------------------------------------
# Discovery & loading
# ---------------------------------------------------------------------------


def find_tasks_dir(takeout_dir: Path) -> Path | None:
    """Locate the Tasks directory anywhere inside a Takeout tree."""
    # Common layout: Takeout/Tasks/ or Takeout/<account>/Tasks/
    children = list(takeout_dir.iterdir()) if takeout_dir.is_dir() else []
    for candidate in [takeout_dir, *children]:
        if not isinstance(candidate, Path):
            continue
        if candidate.is_dir() and (candidate / "Tasks").is_dir():
            return candidate / "Tasks"
    # Fallback: walk the tree
    for root, dirs, files in os.walk(takeout_dir):
        if Path(root).name == "Tasks":
            if any(f.lower().endswith(".json") for f in files):
                return Path(root)
    return None


def load_task_lists(tasks_dir: Path) -> list[dict]:
    """Load and normalise all task-list JSON files from *tasks_dir*."""
    task_lists: list[dict] = []
    for path in sorted(tasks_dir.iterdir()):
        if path.suffix.lower() != ".json":
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue
        if not isinstance(data, dict):
            continue

        # Current Takeout format: one file with kind="tasks#taskLists" that
        # wraps all task lists as top-level items, each with its own items.
        if data.get("kind") == "tasks#taskLists":
            for task_list in data.get("items") or []:
                if isinstance(task_list, dict):
                    task_list.setdefault("title", path.stem)
                    task_list.setdefault("items", [])
                    task_lists.append(task_list)
        else:
            # Older / alternative format: one file per task list.
            data.setdefault("title", path.stem)
            data.setdefault("items", [])
            task_lists.append(data)

    return task_lists


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------


def _parse_date(s: str | None) -> dt.date | None:
    if not s:
        return None
    try:
        return dt.date.fromisoformat(s.split("T")[0])
    except (ValueError, AttributeError):
        return None


def render_tasks_docx(task_lists: list[dict]) -> io.BytesIO:
    """Render *task_lists* to a DOCX document returned as a BytesIO buffer."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    MUTED = RGBColor(0x6B, 0x72, 0x80)
    STRIKE_COLOR = RGBColor(0x9C, 0xA3, 0xAF)

    doc = Document()

    def add_page_break() -> None:
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def add_section_label(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text.upper())
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = MUTED

    def add_task(task: dict, level: int = 0) -> None:
        completed = task.get("status") == "completed"
        title = task.get("title") or "(untitled)"
        checkbox = "☑" if completed else "☐"

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(level * 24)
        p.paragraph_format.space_after = Pt(2)

        title_run = p.add_run(f"{checkbox}  {title}")
        if completed:
            title_run.font.strike = True
            title_run.font.color.rgb = STRIKE_COLOR

        due = _parse_date(task.get("due"))
        if due:
            r = p.add_run(f"   ·   Due: {due.strftime('%b %d, %Y')}")
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED

        comp_date = _parse_date(task.get("completed"))
        if completed and comp_date:
            r = p.add_run(f"  (done {comp_date.strftime('%b %d, %Y')})")
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED

        notes = (task.get("notes") or "").strip()
        if notes:
            np = doc.add_paragraph()
            np.paragraph_format.left_indent = Pt(level * 24 + 20)
            np.paragraph_format.space_after = Pt(4)
            nr = np.add_run(notes)
            nr.font.size = Pt(9)
            nr.font.italic = True
            nr.font.color.rgb = MUTED

        for child in child_map.get(task.get("id", ""), []):
            add_task(child, level + 1)

    # ── Document title ───────────────────────────────────────────────────────
    doc.add_heading("Google Tasks Archive", 0)
    meta = doc.add_paragraph(f"Exported: {dt.date.today().strftime('%B %d, %Y')}")
    for run in meta.runs:
        run.font.color.rgb = MUTED
        run.font.size = Pt(10)

    for task_list in task_lists:
        list_title: str = task_list.get("title") or "Tasks"
        items: list[dict] = task_list.get("items") or []

        doc.add_heading(list_title, 1)

        if not items:
            p = doc.add_paragraph("(no tasks)")
            for run in p.runs:
                run.font.color.rgb = MUTED
                run.font.italic = True
            continue

        items = sorted(items, key=lambda t: t.get("created", ""))

        child_map: dict[str, list[dict]] = {}
        for t in items:
            if t.get("parent"):
                child_map.setdefault(t["parent"], []).append(t)

        root_tasks = [t for t in items if not t.get("parent")]
        active = [t for t in root_tasks if t.get("status") != "completed"]
        completed_tasks = [t for t in root_tasks if t.get("status") == "completed"]

        # ── Active tasks ─────────────────────────────────────────────────────
        add_section_label(f"Active — {len(active)}")
        if active:
            for task in active:
                add_task(task)
        else:
            p = doc.add_paragraph("(all tasks completed)")
            for run in p.runs:
                run.font.color.rgb = MUTED
                run.font.italic = True

        # ── Page break then completed tasks ──────────────────────────────────
        add_page_break()
        add_section_label(f"Completed — {len(completed_tasks)}")
        for task in completed_tasks:
            add_task(task)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
